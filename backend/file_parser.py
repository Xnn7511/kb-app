"""
高分子材料原料知识库 - 增强版文件解析器
支持 PDF(含OCR), DOCX, TXT, CSV, XLSX
"""
import hashlib
import re
from pathlib import Path
from typing import Tuple, Optional, Dict


def compute_checksum(filepath: str) -> str:
    h = hashlib.md5()
    with open(filepath, 'rb') as f:
        for chunk in iter(lambda: f.read(8192), b''):
            h.update(chunk)
    return h.hexdigest()


def _decode_unicode_escapes(text: str) -> str:
    """将 /UNIxxxx 序列解码为实际 Unicode 字符"""
    if not text: return text
    def _repl(m):
        try:
            return chr(int(m.group(1), 16))
        except ValueError:
            return m.group(0)
    # 匹配 /UNIxxxx 或 /UNI-xxxxx 格式（4-5位十六进制）
    text = re.sub(r'/UNI-?([0-9A-Fa-f]{4,5})', _repl, text)
    return text


# 英文关键词到中文的映射（用于标题/摘要翻译）
EN_ZH_MAP = {
    'anti-block': '抗粘', 'anti block': '抗粘', 'antiblock': '抗粘',
    'slip': '爽滑', 'masterbatch': '母料', 'additive': '添加剂',
    'polypropylene': '聚丙烯', 'polyethylene': '聚乙烯', 'polymer': '聚合物',
    'resin': '树脂', 'film': '薄膜', 'bopp': 'BOPP', 'cpp': 'CPP',
    'density': '密度', 'melt flow': '熔融指数', 'mfr': '熔融指数',
    'melting point': '熔点', 'tensile': '拉伸', 'strength': '强度',
    'haze': '雾度', 'gloss': '光泽度', 'vicat': '维卡软化点',
    'adhesive': '胶粘剂', 'primer': '底涂', 'coating': '涂层',
    'flame retardant': '阻燃', 'antistatic': '抗静电',
    'antifog': '防雾', 'matte': '哑光', 'matting': '消光',
    'pearl': '珠光', 'whitening': '增白', 'white': '增白',
    'technical data sheet': '技术数据表', 'tds': '技术数据表',
    'product description': '产品描述', 'application': '应用',
    'properties': '性能', 'specification': '规格',
    'general properties': '一般性能', 'general': '一般',
    'evolue': 'Evolue', 'prime polymer': 'Prime聚合物',
    'plastomer': '塑性体', 'elastomer': '弹性体',
    'ethylene': '乙烯', 'butene': '丁烯', 'copolymer': '共聚物',
    'polyolefin': '聚烯烃', 'olefin': '烯烃',
    'compatibilizer': '相容剂', 'coupling': '偶联',
    'graft': '接枝', 'modified': '改性',
    'high density': '高密度', 'low density': '低密度',
    'linear': '线性', 'metallocene': '茂金属',
    'catalyst': '催化剂', 'reaction': '反应',
    'extrusion': '挤出', 'injection': '注塑',
    'blow molding': '吹塑', 'casting': '流延',
    'thermal': '热', 'stability': '稳定性',
    'weathering': '耐候', 'uv': '抗UV', 'ultraviolet': '紫外',
    'transparent': '透明', 'transparency': '透明度',
    'clear': '透明', 'opacity': '不透明度',
    'food contact': '食品接触', 'fda': 'FDA',
    'rohs': 'RoHS', 'reach': 'REACH',
    'safety data sheet': '安全数据表', 'sds': '安全数据表', 'msds': '安全数据表',
    'flash point': '闪点', 'toxicity': '毒性',
    'storage': '储存', 'handling': '操作',
    'disclaimer': '免责声明', 'warranty': '质保',
    'revision': '修订', 'date': '日期',
}


def translate_en_title(title: str) -> str:
    """将英文标题翻译为中文（基于关键词映射）"""
    if not title: return title
    t = title.lower()
    result = title
    # 按关键词长度降序替换，避免短词覆盖长词
    for en, zh in sorted(EN_ZH_MAP.items(), key=lambda x: -len(x[0])):
        if en in t:
            # 保留原大小写风格，替换匹配部分
            result = re.sub(r'(?i)' + re.escape(en), zh, result)
    return result


def _clean_text(text: str) -> str:
    if not text: return ""
    # 先解码 /UNIxxxx 序列
    text = _decode_unicode_escapes(text)
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    for a, b in [('\ufb01','fi'),('\ufb02','fl'),('\ufb00','ff'),('\ufb03','ffi'),('\ufb04','ffl'),
                 ('\u2013','-'),('\u2014','--'),('\u2018',"'"),('\u2019',"'"),
                 ('\u201c','"'),('\u201d','"'),('\u2026','...'),('\u00a0',' ')]:
        text = text.replace(a, b)
    return text.strip()


def parse_pdf(filepath: str) -> str:
    text = ""
    # 策略1: pypdf
    try:
        import pypdf
        reader = pypdf.PdfReader(filepath)
        parts = []
        for i, page in enumerate(reader.pages):
            pt = page.extract_text()
            if pt and len(pt.strip()) > 5:
                parts.append(f"[第{i+1}页]\n{pt}")
        text = "\n\n".join(parts) if parts else ""
    except Exception:
        pass

    # 策略2: pypdfium2
    if len(text) < 100:
        try:
            import pypdfium2 as pdfium
            pdf = pdfium.PdfDocument(filepath)
            parts = []
            for i in range(len(pdf)):
                pt = pdf[i].get_text()
                if pt and len(pt.strip()) > 5:
                    parts.append(f"[第{i+1}页]\n{pt}")
            pdf.close()
            if parts: text = "\n\n".join(parts)
        except Exception:
            pass

    # 策略3: OCR（图片型PDF）
    if len(text) < 100:
        try:
            from pdf2image import convert_from_path
            import pytesseract
            # 限制OCR页数提高速度
            total_pages = _get_pdf_page_count(filepath)
            ocr_pages = min(total_pages, 8)
            images = convert_from_path(filepath, first_page=ocr_pages, dpi=200)
            parts = []
            for i, img in enumerate(images):
                pt = pytesseract.image_to_string(img, lang='chi_sim+eng')
                if pt and len(pt.strip()) > 5:
                    parts.append(f"[第{i+1}页(OCR)]\n{pt}")
            if parts: text = "\n\n".join(parts)
        except Exception:
            pass

    if not text:
        text = f"[此PDF为图片格式，共{_get_pdf_page_count(filepath)}页。文本提取有限，建议查看原始文件。]"
    return _clean_text(text)


def _get_pdf_page_count(filepath: str) -> int:
    try:
        import pypdf; return len(pypdf.PdfReader(filepath).pages)
    except: return 0


def parse_docx(filepath: str) -> str:
    try:
        from docx import Document
        doc = Document(filepath)
        parts = []
        for p in doc.paragraphs:
            if p.text.strip(): parts.append(p.text)
        for ti, t in enumerate(doc.tables):
            parts.append(f"\n[表格 {ti+1}]")
            for row in t.rows:
                parts.append(" | ".join(c.text.strip() for c in row.cells))
        return _clean_text("\n".join(parts))
    except Exception as e:
        return f"[DOCX 解析错误: {str(e)}]"


def parse_txt(filepath: str) -> str:
    for enc in ['utf-8', 'gbk', 'gb2312', 'gb18030', 'big5', 'latin-1']:
        try:
            with open(filepath, 'r', encoding=enc) as f:
                text = f.read()
                if _is_valid_text(text): return _clean_text(text)
        except: continue
    with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
        return _clean_text(f.read())


def _is_valid_text(text: str) -> bool:
    return text and text.count('\ufffd') < len(text) * 0.01


def parse_csv(filepath: str) -> str:
    import csv
    for enc in ['utf-8', 'gbk', 'gb2312', 'gb18030']:
        try:
            parts = []
            with open(filepath, 'r', encoding=enc, newline='') as f:
                for row in csv.reader(f):
                    if any(c.strip() for c in row):
                        parts.append(" | ".join(row))
            return _clean_text("\n".join(parts))
        except: continue
    return "[CSV 解析失败]"


def parse_xlsx(filepath: str) -> str:
    try:
        import openpyxl
        wb = openpyxl.load_workbook(filepath, data_only=True)
        parts = []
        for sn in wb.sheetnames:
            ws = wb[sn]
            parts.append(f"\n[工作表: {sn}]")
            for row in ws.iter_rows(values_only=True):
                if any(c is not None for c in row):
                    parts.append(" | ".join(str(c) if c is not None else "" for c in row))
        wb.close()
        return _clean_text("\n".join(parts))
    except Exception as e:
        return f"[XLSX 解析错误: {str(e)}]"


PARSERS = {'.pdf': parse_pdf, '.docx': parse_docx, '.doc': parse_docx,
           '.txt': parse_txt, '.md': parse_txt, '.csv': parse_csv,
           '.xlsx': parse_xlsx, '.xls': parse_xlsx}


def parse_file(filepath: str) -> Tuple[str, str]:
    path = Path(filepath)
    ext = path.suffix.lower()
    parser = PARSERS.get(ext)
    return (parser(filepath), ext) if parser else (f"[不支持: {ext}]", ext)


def detect_language(text: str) -> str:
    if not text: return 'en'
    cn = len(re.findall(r'[\u4e00-\u9fff]', text))
    en = len(re.findall(r'[a-zA-Z]{3,}', text))
    total = max(cn + en, 1)
    return 'zh' if cn / total > 0.3 else ('mixed' if en / total < 0.7 else 'en')


def extract_tds_data(text: str) -> Dict:
    """从文本提取TDS物性数据"""
    data = dict.fromkeys(['mfr','density','tensile_strength','elongation',
        'flexural_modulus','impact_strength','hdt','vicat','melting_point',
        'haze','gloss','ash_content'])
    patterns = {
        'mfr': [r'MFR?\s*[：:]\s*([\d.]+)',r'MFR?\s*\(?\d+°?C.*?\)?\s*[：:]\s*([\d.]+)',
                r'Melt\s*Flow\s*(?:Rate|Index).*?[：:]\s*([\d.]+)',
                r'熔(?:融|体).*?(?:指数|速率).*?[：:]\s*([\d.]+)',
                r'(?:MFR|MI)\s+([\d.]+)\s*g',r'MFR\s+([\d.]+)',
                r'熔体指数.*?g/10min\s+([\d.]+)',r'熔体指数.*?g/10min\s*\n\s*([\d.]+)'],
        'density': [r'[Dd]ensity.*?[：:]\s*([\d.]+)\s*g/cm',r'密度\s*[：:]\s*([\d.]+)\s*g/cm',
                    r'密度\s+([\d.]+)\s*g/cm',r'Density\s+([\d.]+)\s*g/cm',
                    r'密度.*?([\d.]+)\s*g/cm',r'density.*?([\d.]+)\s*g/cm',
                    r'密度.*?g/cm3?\s+([\d.]+)',r'密度.*?g/cm3?\s*\n\s*([\d.]+)',
                    # kg/m³ 需要除以1000转换
                    r'密度\s*[：:]\s*([\d.]+)\s*kg/m',r'Density\s*[：:]\s*([\d.]+)\s*kg/m'],
        'tensile_strength': [r'拉伸强度\s*[：:]\s*([\d.]+)',r'Tensile\s*(?:Strength|stress).*?[：:]\s*([\d.]+)',
                             r'拉伸.*?强度\s+([\d.]+)\s*MPa'],
        'melting_point': [r'熔(?:融)?点\s*[：:]\s*([\d.]+)\s*[°℃]',r'Melting\s*(?:Point|Temperature).*?[：:]\s*([\d.]+)',
                          r'(?:Tm|mp)\s*[：:]\s*([\d.]+)',r'熔点\s+([\d.]+)\s*[°℃]'],
        'haze': [r'雾度\s*[：:]\s*([\d.]+)',r'Haze\s*[：:]\s*([\d.]+)',r'雾度\s+([\d.]+)\s*%',r'Haze\s+([\d.]+)\s*%'],
        'hdt': [r'热变形温度\s*[：:]\s*([\d.]+)',r'HDT\s*[：:]\s*([\d.]+)',r'HDT\s+([\d.]+)\s*[°℃]',r'热变形\s+([\d.]+)\s*[°℃]'],
        'vicat': [r'维卡\s*[：:]\s*([\d.]+)',r'Vicat\s*[：:]\s*([\d.]+)',r'维卡.*?([\d.]+)\s*[°℃]'],
        'gloss': [r'光泽度\s*[：:]\s*([\d.]+)',r'Gloss\s*[：:]\s*([\d.]+)',r'光泽.*?([\d.]+)\s*%'],
        'ash_content': [r'灰分\s*[：:]\s*([\d.]+)',r'Ash\s*[：:]\s*([\d.]+)',r'灰分\s+([\d.]+)\s*%'],
    }
    for key, pats in patterns.items():
        for pat in pats:
            m = re.search(pat, text, re.IGNORECASE)
            if m:
                try:
                    val = float(m.group(1))
                    # 如果是密度且匹配到 kg/m³ 模式，需要除以 1000
                    if key == 'density' and 'kg/m' in pat:
                        val = val / 1000.0
                    # 密度合理范围 0.5~3.0 g/cm³；其他参数范围更宽
                    if key == 'density' and not (0.3 < val < 3.5):
                        continue
                    if 0.001 < val < 10000:
                        data[key] = round(val, 3)
                        break
                except: continue
    return data


def extract_manufacturer(title: str, text: str) -> str:
    """从标题和文本中提取厂家信息"""
    combined = (title + ' ' + text[:1000]).lower()

    manufacturers = {
        '高博尔': ['高博尔', 'gaoboer', 'global'],
        '广东基烁': ['基烁', 'jishuo'],
        '康斯坦普': ['constab', '康斯坦普'],
        '埃克森美孚': ['exxon', '埃克森', 'vistamaxx'],
        '博禄': ['borouge', '博禄'],
        '巴赛尔/利安德巴赛尔': ['lyondellbasell', '巴赛尔', 'basell', 'polybatch'],
        '三井化学/TAFMER': ['tafmer', '三井', 'mitsui'],
        'SK致新/SKGC': ['sk geo', 'skgc', 'sk global', 'yuplene'],
        '陶氏/DOW': ['dow', '陶氏', 'bynel'],
        '阿科玛/SKFP': ['orevac', 'skfp', 'arkema', '阿科玛'],
        'Prime Evolue': ['prime evolue', 'evolue', 'prime polymer'],
        'Polymateria': ['polymateria'],
        '叶心石化': ['叶心', 'yexin'],
        '佛山塑兴': ['塑兴', 'suxing', '佛山'],
        '福融': ['福融', 'furong'],
        '鸿盛': ['鸿盛', 'hongsheng'],
        'CPJJ': ['cpjj'],
        '中国石化/燕山': ['燕山', '中国石化', '中石化', 'sinopec', '燕化'],
        '链行走': ['链行走', 'lianxingzou', 'chain-walking', 'chain walking', 'walking polymer'],
    }

    for mfr, keywords in manufacturers.items():
        for kw in keywords:
            if kw in combined:
                return mfr
    return ''


def generate_smart_summary(text: str, tds_data: Dict, max_len: int = 120) -> str:
    """生成智能摘要，优先展示关键产品指标"""
    if not text: return ''

    parts = []

    # 优先展示TDS关键指标
    tds_priority = ['melting_point', 'mfr', 'density', 'hdt', 'vicat', 'haze', 'tensile_strength', 'gloss']
    tds_labels = {'melting_point':'熔点','mfr':'MFR','density':'密度','hdt':'热变形温度',
                  'vicat':'维卡软化点','haze':'雾度','tensile_strength':'拉伸强度','gloss':'光泽度'}
    tds_units = {'melting_point':'°C','mfr':'g/10min','density':'g/cm³','hdt':'°C',
                 'vicat':'°C','haze':'%','tensile_strength':'MPa','gloss':'%'}

    for key in tds_priority:
        val = tds_data.get(key)
        if val is not None:
            parts.append(f"{tds_labels.get(key, key)}: {val}{tds_units.get(key, '')}")

    # 如果TDS数据不足，用原文前80字符补充
    if not parts:
        # 尝试找第一句有意义的中文
        m = re.search(r'[\u4e00-\u9fff][\u4e00-\u9fff，。、；：！？\w\s]{20,80}', text)
        if m:
            parts.append(m.group().strip())
        else:
            parts.append(text[:80].strip())

    summary = ' | '.join(parts)
    if len(summary) > max_len:
        summary = summary[:max_len-3] + '...'
    return summary
